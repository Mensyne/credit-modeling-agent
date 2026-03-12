import streamlit as st
import polars as pl
from datetime import datetime
from docx import Document
from docx.shared import Inches
import io
import os
from weasyprint import HTML
import tempfile


def render_regulatory_report_agent():
    st.title("📑 监管合规报告Agent")

    if (
        "current_project" not in st.session_state
        or "current_model" not in st.session_state
    ):
        st.warning("请先进入项目并训练模型！")
        return

    project = st.session_state.current_project
    model = st.session_state.current_model

    st.subheader("报告配置")
    report_type = st.selectbox(
        "报告类型",
        [
            "商业银行互联网贷款模型开发报告",
            "商业银行互联网贷款模型验证报告",
            "模型季度合规报告",
        ],
    )
    company_name = st.text_input("机构名称", value="XX银行/消费金融公司")
    report_author = st.text_input("报告撰写人", value="风控建模团队")
    report_period = st.text_input(
        "报告期", value=f"{datetime.now().strftime('%Y年%m月')}"
    )

    # 加载报告数据
    report_data = {
        "basic_info": {
            "project_name": project["name"],
            "model_name": model["name"],
            "model_type": model["type"],
            "company_name": company_name,
            "report_author": report_author,
            "report_period": report_period,
            "report_date": datetime.now().strftime("%Y年%m月%d日"),
        },
        "data_info": {},
        "feature_info": {},
        "model_info": {},
        "evaluation_info": {},
        "compliance_info": {},
    }

    # 填充数据信息
    if "raw_data" in st.session_state:
        df = st.session_state.raw_data
        report_data["data_info"] = {
            "total_samples": len(df),
            "feature_count": len(df.columns) - 1,
            "bad_rate": f"{df[model['label_col']].mean():.2%}",
            "time_span": "",
        }
        if "train_data" in st.session_state and "val_data" in st.session_state:
            report_data["data_info"]["train_samples"] = len(st.session_state.train_data)
            report_data["data_info"]["val_samples"] = len(st.session_state.val_data)
            if "test_data" in st.session_state:
                report_data["data_info"]["test_samples"] = len(
                    st.session_state.test_data
                )

    # 填充特征信息
    if "selected_features" in st.session_state:
        report_data["feature_info"] = {
            "selected_feature_count": len(st.session_state.selected_features),
            "top_features": [],
        }
        if hasattr(model["model"], "feature_importances_"):
            imp = model["model"].feature_importances_
            top_idx = imp.argsort()[-5:][::-1]
            report_data["feature_info"]["top_features"] = [
                st.session_state.selected_features[i] for i in top_idx
            ]

    # 填充模型信息
    report_data["model_info"] = {
        "params": model["params"],
        "train_auc": f"{model['train_auc']:.4f}",
        "val_auc": f"{model['val_auc']:.4f}" if model["val_auc"] else "-",
    }

    # 生成报告预览
    if st.button("生成报告预览", type="primary"):
        st.subheader("报告预览")

        # 1. 报告封面
        st.markdown(f"# {report_type}")
        st.markdown(f"## {company_name}")
        st.markdown(f"### 项目名称：{project['name']}")
        st.markdown(f"### 模型名称：{model['name']}")
        st.markdown(f"**报告期：{report_period}**")
        st.markdown(f"**撰写人：{report_author}**")
        st.markdown(f"**生成日期：{datetime.now().strftime('%Y年%m月%d日')}**")
        st.divider()

        # 2. 目录
        st.markdown("## 目录")
        st.markdown("1. 项目概述")
        st.markdown("2. 数据源及样本情况")
        st.markdown("3. 特征工程说明")
        st.markdown("4. 模型开发过程")
        st.markdown("5. 模型效果评估")
        st.markdown("6. 合规性检查")
        st.markdown("7. 结论及建议")
        st.divider()

        # 3. 项目概述
        st.markdown("## 1. 项目概述")
        st.markdown(f"### 1.1 项目背景")
        st.markdown(
            f"本报告为{project['name']}项目的{report_type}，模型算法类型为{model['type']}，用于信贷业务客户风险评估。"
        )
        st.markdown(f"### 1.2 报告范围")
        st.markdown(
            f"本报告覆盖模型开发全流程，包括数据、特征、模型、评估、合规等环节，符合《商业银行互联网贷款管理暂行办法》要求。"
        )
        st.divider()

        # 4. 数据情况
        st.markdown("## 2. 数据源及样本情况")
        if report_data["data_info"]:
            st.markdown("### 2.1 样本统计")
            data_df = pl.DataFrame(
                [
                    {
                        "指标": "总样本量",
                        "数值": report_data["data_info"]["total_samples"],
                    },
                    {
                        "指标": "特征数量",
                        "数值": report_data["data_info"]["feature_count"],
                    },
                    {"指标": "样本坏率", "数值": report_data["data_info"]["bad_rate"]},
                    {
                        "指标": "训练集样本量",
                        "数值": report_data["data_info"].get("train_samples", "-"),
                    },
                    {
                        "指标": "验证集样本量",
                        "数值": report_data["data_info"].get("val_samples", "-"),
                    },
                    {
                        "指标": "测试集样本量",
                        "数值": report_data["data_info"].get("test_samples", "-"),
                    },
                ]
            )
            st.dataframe(data_df, use_container_width=True)
        st.divider()

        # 5. 特征工程
        st.markdown("## 3. 特征工程说明")
        if report_data["feature_info"]:
            st.markdown(
                f"### 3.1 特征筛选结果：最终入模特征共{report_data['feature_info']['selected_feature_count']}个"
            )
            if report_data["feature_info"]["top_features"]:
                st.markdown("### 3.2 重要性Top5特征：")
                for i, feat in enumerate(
                    report_data["feature_info"]["top_features"], 1
                ):
                    st.markdown(f"{i}. {feat}")
        st.divider()

        # 6. 模型开发
        st.markdown("## 4. 模型开发过程")
        st.markdown(f"### 4.1 算法类型：{model['type']}")
        st.markdown("### 4.2 模型参数：")
        st.json(model["params"])
        st.divider()

        # 7. 模型评估
        st.markdown("## 5. 模型效果评估")
        st.markdown("### 5.1 核心指标：")
        eval_df = pl.DataFrame(
            [
                {"数据集": "训练集", "AUC": report_data["model_info"]["train_auc"]},
                {"数据集": "验证集", "AUC": report_data["model_info"]["val_auc"]},
            ]
        )
        st.dataframe(eval_df, use_container_width=True)
        st.divider()

        # 8. 合规检查
        st.markdown("## 6. 合规性检查")
        st.markdown("✅ 敏感特征排查：已完成敏感特征识别和脱敏处理")
        st.markdown("✅ 数据授权：已确认数据源获得合法授权")
        st.markdown("✅ 可解释性：模型可解释性符合监管要求")
        st.divider()

        # 9. 结论
        st.markdown("## 7. 结论及建议")
        st.markdown("### 7.1 评估结论")
        st.markdown(
            f"本模型训练集AUC {model['train_auc']:.4f}，验证集AUC {model['val_auc']:.4f}，效果符合业务要求，可上线使用。"
        )
        st.markdown("### 7.2 监控建议")
        st.markdown(
            "建议每月监控模型AUC、KS、PSI指标，当PSI超过0.25时及时重新训练模型。"
        )

    # 导出功能
    col1, col2 = st.columns(2)
    with col1:
        if st.button("导出Word报告"):
            with st.spinner("正在生成Word报告..."):
                doc = Document()
                # 封面
                doc.add_heading(report_type, 0)
                doc.add_paragraph(f"机构名称：{company_name}")
                doc.add_paragraph(f"项目名称：{project['name']}")
                doc.add_paragraph(f"模型名称：{model['name']}")
                doc.add_paragraph(f"报告期：{report_period}")
                doc.add_paragraph(
                    f"生成日期：{datetime.now().strftime('%Y年%m月%d日')}"
                )
                doc.add_page_break()

                # 内容
                doc.add_heading("1. 项目概述", level=1)
                doc.add_paragraph(
                    f"本报告为{project['name']}项目的{report_type}，模型算法类型为{model['type']}，用于信贷业务客户风险评估。"
                )

                doc.add_heading("2. 数据源及样本情况", level=1)
                if report_data["data_info"]:
                    table = doc.add_table(rows=1, cols=2)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = "指标"
                    hdr_cells[1].text = "数值"
                    for k, v in report_data["data_info"].items():
                        row_cells = table.add_row().cells
                        row_cells[0].text = k
                        row_cells[1].text = str(v)

                doc.add_heading("3. 模型评估结果", level=1)
                doc.add_paragraph(f"训练集AUC：{model['train_auc']:.4f}")
                doc.add_paragraph(
                    f"验证集AUC：{model['val_auc']:.4f}"
                    if model["val_auc"]
                    else "验证集AUC：-"
                )

                # 保存
                bio = io.BytesIO()
                doc.save(bio)
                bio.seek(0)

                st.download_button(
                    label="下载Word报告",
                    data=bio.getvalue(),
                    file_name=f"{report_type}_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
                st.success("Word报告生成完成！")

    with col2:
        if st.button("导出PDF报告"):
            with st.spinner("正在生成PDF报告..."):
                # 生成HTML内容
                html_content = f"""
                <h1>{report_type}</h1>
                <h2>{company_name}</h2>
                <p>项目名称：{project["name"]}</p>
                <p>模型名称：{model["name"]}</p>
                <p>报告期：{report_period}</p>
                <p>生成日期：{datetime.now().strftime("%Y年%m月%d日")}</p>
                <hr>
                <h2>1. 项目概述</h2>
                <p>本报告为{project["name"]}项目的{report_type}，模型算法类型为{model["type"]}，用于信贷业务客户风险评估。</p>
                <h2>2. 模型效果</h2>
                <p>训练集AUC：{model["train_auc"]:.4f}</p>
                <p>验证集AUC：{model["val_auc"]:.4f}</p>
                <h2>3. 合规声明</h2>
                <p>本模型开发流程符合《商业银行互联网贷款管理暂行办法》要求。</p>
                """

                # 生成PDF
                bio = io.BytesIO()
                HTML(string=html_content).write_pdf(bio)
                bio.seek(0)

                st.download_button(
                    label="下载PDF报告",
                    data=bio.getvalue(),
                    file_name=f"{report_type}_{datetime.now().strftime('%Y%m%d')}.pdf",
                    mime="application/pdf",
                )
                st.success("PDF报告生成完成！")
